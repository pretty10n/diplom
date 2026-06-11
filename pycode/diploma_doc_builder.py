"""Сборщик Word-документов ВКР по шаблону ВГТУ (бакалавры ИБ)."""

from __future__ import annotations

from pathlib import Path
from typing import Iterable, Sequence

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor
from docx.table import Table

# Шаблон ВКР (бакалавры ИБ): Times New Roman 14 pt, поля 30/15/20/20 мм
FONT_NAME = "Times New Roman"
BODY_SIZE = Pt(14)
HEADING_SIZE = Pt(14)
LINE_SPACING = 1.5
FIRST_LINE_INDENT = Cm(1.25)
RED = RGBColor(0xFF, 0x00, 0x00)

MARGIN_LEFT = Mm(30)
MARGIN_RIGHT = Mm(15)
MARGIN_TOP = Mm(20)
MARGIN_BOTTOM = Mm(20)


def _set_run_font(run, size: Pt = BODY_SIZE, bold: bool = False, color=None) -> None:
    run.font.name = FONT_NAME
    run.font.size = size
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    r = run._element
    r_pr = r.get_or_add_rPr()
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), FONT_NAME)
    r_fonts.set(qn("w:hAnsi"), FONT_NAME)
    r_fonts.set(qn("w:cs"), FONT_NAME)
    r_pr.insert(0, r_fonts)


def _paragraph_format(
    paragraph,
    *,
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    indent: float | None = FIRST_LINE_INDENT,
    size: Pt = BODY_SIZE,
    bold: bool = False,
    color=None,
) -> None:
    fmt = paragraph.paragraph_format
    fmt.alignment = align
    fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    fmt.line_spacing = LINE_SPACING
    fmt.first_line_indent = indent
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    if paragraph.runs:
        for run in paragraph.runs:
            _set_run_font(run, size=size, bold=bold, color=color)
    else:
        run = paragraph.add_run()
        _set_run_font(run, size=size, bold=bold, color=color)


def _add_page_number_run(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_begin)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    run._r.append(instr)
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    run._r.append(fld_sep)
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_end)


def setup_document() -> Document:
    doc = Document()
    for section in doc.sections:
        section.page_height = Mm(297)
        section.page_width = Mm(210)
        section.left_margin = MARGIN_LEFT
        section.right_margin = MARGIN_RIGHT
        section.top_margin = MARGIN_TOP
        section.bottom_margin = MARGIN_BOTTOM
        footer = section.footer
        footer.is_linked_to_previous = False
        if footer.paragraphs:
            _add_page_number_run(footer.paragraphs[0])
        else:
            _add_page_number_run(footer.add_paragraph())

    normal = doc.styles["Normal"]
    normal.font.name = FONT_NAME
    normal.font.size = BODY_SIZE
    pf = normal.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = LINE_SPACING
    pf.first_line_indent = FIRST_LINE_INDENT

    for style_name in ("Heading 1", "Heading 2"):
        try:
            style = doc.styles[style_name]
            style.font.name = FONT_NAME
            style.font.size = HEADING_SIZE
            style.font.bold = True
            spf = style.paragraph_format
            spf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            spf.line_spacing = LINE_SPACING
        except KeyError:
            pass
    return doc


def add_chapter_title(doc: Document, text: str) -> None:
    """Заголовок раздела (главы): 14 pt, полужирный, по ширине с абзацным отступом."""
    doc.add_page_break()
    try:
        p = doc.add_paragraph(text, style="Heading 1")
    except KeyError:
        p = doc.add_paragraph()
        p.add_run(text)
    _paragraph_format(
        p,
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
        indent=FIRST_LINE_INDENT,
        size=HEADING_SIZE,
        bold=True,
    )


def add_structural_title(doc: Document, text: str, *, page_break: bool = True) -> None:
    """Структурный элемент: ВВЕДЕНИЕ, СОДЕРЖАНИЕ, ЗАКЛЮЧЕНИЕ — по центру, 14 pt, полужирный."""
    if page_break:
        doc.add_page_break()
    try:
        p = doc.add_paragraph(text.upper(), style="Heading 1")
    except KeyError:
        p = doc.add_paragraph()
        p.add_run(text.upper())
    _paragraph_format(
        p,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        indent=None,
        size=HEADING_SIZE,
        bold=True,
    )


def add_section_heading(doc: Document, text: str) -> None:
    try:
        p = doc.add_paragraph(text, style="Heading 2")
    except KeyError:
        p = doc.add_paragraph()
        p.add_run(text)
    _paragraph_format(
        p,
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
        indent=None,
        size=HEADING_SIZE,
        bold=True,
    )


def add_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    _paragraph_format(p)


def add_bullet_list(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.text = item
        _paragraph_format(p, indent=None)


def add_table(doc: Document, headers: Sequence[str], rows: Sequence[Sequence[str]]) -> Table:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for paragraph in hdr_cells[i].paragraphs:
            _paragraph_format(paragraph, indent=None, bold=False)
    for r_idx, row in enumerate(rows, start=1):
        for c_idx, value in enumerate(row):
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = str(value)
            for paragraph in cell.paragraphs:
                _paragraph_format(paragraph, indent=None)
    doc.add_paragraph()
    return table


def add_table_caption(doc: Document, caption: str) -> None:
    p = doc.add_paragraph(caption)
    _paragraph_format(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent=FIRST_LINE_INDENT)


def add_figure_image(doc: Document, image_path: Path, caption: str) -> None:
    if image_path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = None
        run = p.add_run()
        run.add_picture(str(image_path), width=Cm(15))
    else:
        add_figure_placeholder(doc, f"СХЕМА: {image_path.name}", caption)
        return
    cap = doc.add_paragraph(caption)
    _paragraph_format(cap, align=WD_ALIGN_PARAGRAPH.CENTER, indent=None)


def add_figure_placeholder(doc: Document, placeholder: str, caption: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(f"({placeholder})")
    _paragraph_format(p, align=WD_ALIGN_PARAGRAPH.CENTER, indent=None, color=RED)
    p.runs[0].font.color.rgb = RED
    cap = doc.add_paragraph(caption)
    _paragraph_format(cap, align=WD_ALIGN_PARAGRAPH.CENTER, indent=None)


def add_bibliography_entry(doc: Document, number: int, text: str) -> None:
    p = doc.add_paragraph()
    p.add_run(f"{number} {text}")
    fmt = p.paragraph_format
    fmt.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    fmt.line_spacing = LINE_SPACING
    fmt.first_line_indent = Cm(-1.25)
    fmt.left_indent = Cm(1.25)
    for run in p.runs:
        _set_run_font(run)


def save_document(doc: Document, output_path: Path) -> Path:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path

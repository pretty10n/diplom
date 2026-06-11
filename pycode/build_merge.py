"""Сборка полного ВКР: содержание + введение + 3 главы + заключение + литература."""

from __future__ import annotations

import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent))

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docxcompose.composer import Composer

from diploma_doc_builder import HEADING_SIZE, setup_document, save_document, _paragraph_format
from paths import GENERATED_DIR


PARTS = [
    "Введение.docx",
    "Глава_1_Анализ.docx",
    "Глава_2_Методы.docx",
    "Глава_3_Реализация.docx",
    "Заключение.docx",
    "Список_литературы.docx",
]

ALT_PARTS = {
    "Введение.docx": "Введение_new.docx",
    "Глава_1_Анализ.docx": "Глава_1_Анализ_new.docx",
    "Глава_2_Методы.docx": "Глава_2_Методы_new.docx",
    "Глава_3_Реализация.docx": "Глава_3_Реализация_new.docx",
    "Заключение.docx": "Заключение_new.docx",
    "Список_литературы.docx": "Список_литературы_new.docx",
}


def _add_toc_field(paragraph) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    # Уровни 1–2: структурные элементы, главы и подразделы
    instr.text = r'TOC \o "1-2" \h \z \u'
    run._r.append(instr)

    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    run._r.append(fld_sep)

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_end)


def add_contents_page(doc: Document) -> None:
    p_title = doc.add_paragraph()
    p_title.add_run("СОДЕРЖАНИЕ")
    _paragraph_format(
        p_title,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        indent=None,
        size=HEADING_SIZE,
        bold=True,
    )
    p_toc = doc.add_paragraph()
    _paragraph_format(p_toc, indent=None)
    _add_toc_field(p_toc)
    doc.add_page_break()


def _resolve_part(name: str) -> Path:
    primary = GENERATED_DIR / name
    if primary.exists():
        return primary
    alt = GENERATED_DIR / ALT_PARTS.get(name, "")
    if alt and alt.exists():
        return alt
    raise FileNotFoundError(f"Не найден фрагмент: {name}")


def ensure_all_parts() -> list[Path]:
    from build_introduction import build as build_intro
    from build_chapter_1 import build as build1
    from build_chapter_2 import build as build2
    from build_chapter_3 import build as build3
    from build_conclusion import build as build_conclusion
    from build_references import build as build_references

    build_intro()
    build1()
    build2()
    build3()
    build_conclusion()
    build_references()
    return [_resolve_part(name) for name in PARTS]


def build(rebuild_parts: bool = True) -> Path:
    GENERATED_DIR.mkdir(parents=True, exist_ok=True)
    paths = ensure_all_parts() if rebuild_parts else [_resolve_part(n) for n in PARTS]

    master = setup_document()
    add_contents_page(master)

    composer = Composer(master)
    for part_path in paths:
        composer.append(Document(str(part_path)))

    output = GENERATED_DIR / "ВКР_Душина_полная.docx"
    try:
        return save_document(composer.doc, output)
    except PermissionError:
        return save_document(composer.doc, GENERATED_DIR / "ВКР_Душина_полная_new.docx")


if __name__ == "__main__":
    print(build())

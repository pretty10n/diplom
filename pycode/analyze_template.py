"""Анализ шаблона ВКР: стили, поля, абзацы с требованиями."""

from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

path = Path(sys.argv[1]) if len(sys.argv) > 1 else Path("E:/Project-A/generated/vkr_template_ib.docx")
doc = Document(str(path))
sec = doc.sections[0]
print("=== SECTION ===")
print(f"left={sec.left_margin.mm} right={sec.right_margin.mm} top={sec.top_margin.mm} bottom={sec.bottom_margin.mm}")
print(f"page {sec.page_width.mm}x{sec.page_height.mm}")

print("\n=== STYLES (used in doc) ===")
seen = set()
for p in doc.paragraphs:
    name = p.style.name if p.style else "?"
    if name in seen:
        continue
    seen.add(name)
    pf = p.paragraph_format
    align = pf.alignment
    if p.runs:
        r = p.runs[0]
        print(
            f"- {name}: font={r.font.name} size={r.font.size} bold={r.font.bold} "
            f"align={align} indent={pf.first_line_indent} line={pf.line_spacing}"
        )

print("\n=== KEY PARAGRAPHS (first 80 non-empty) ===")
count = 0
for p in doc.paragraphs:
    t = p.text.strip()
    if not t:
        continue
    style = p.style.name if p.style else ""
    print(f"[{style}] {t[:120]}")
    count += 1
    if count >= 80:
        break
